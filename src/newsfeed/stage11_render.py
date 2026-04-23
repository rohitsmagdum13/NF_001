"""Stage 11 — Render assembled newsfeed to .docx and .md.

Inputs
------
    Disk:   outputs/{run_date}/horizon_payload.json  (AssembledNewsfeed)

Outputs
-------
    Disk:   outputs/{run_date}/Water_Newsfeed.docx   (formatted Word document)
    Disk:   outputs/{run_date}/Water_Newsfeed.md     (markdown archive)

What this stage does
--------------------
    Reads the ``AssembledNewsfeed`` written by Stage 9 and produces two files:

    .docx — python-docx document matching the structure of
            ``reference/Water_Newsfeed.md``:
            * Title heading + issue date
            * Numbered index of sections
            * Sections (Heading 2) → Jurisdiction sub-headings (Heading 3)
            * Each draft entry rendered with bold titles, italic source
              attributions, bullet list items, and body paragraphs
            * Upcoming Dates placeholder, Previous Newsfeeds placeholder,
              copyright footer

    .md   — plain markdown archive of the same content (easier to diff/review).

    Each draft entry's ``draft_text`` is parsed line-by-line:
        ``**TITLE**``        → bold paragraph
        ``*(Source: ...)*``  → italic paragraph
        ``- ITEM``           → bulleted list item
        everything else      → normal paragraph

Run directly
------------
    uv run python -m newsfeed.stage11_render
    uv run python -m newsfeed.stage11_render --date 2026-04-22
"""

from __future__ import annotations

import re
import sys
from datetime import date
from pathlib import Path

from docx import Document as _make_docx  # noqa: N813 — docx.Document is a factory fn, not a class
from docx.document import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from loguru import logger

from newsfeed.config import get_settings
from newsfeed.schemas import AssembledEntry, AssembledNewsfeed

_NEWSFEED_TITLE = "Water Newsfeed"
_LEGISLATION_HOTLINE_BODY = (
    "The Legislation Hotline is a value-added service available to our clients that can "
    "assist with your legislation-related questions. For example, you may want to know what "
    "stage a Bill is at, whether an Act has received Assent, or when certain legislation is "
    "commencing.\n\n"
    "Just a phone call away and covering all Australian jurisdictions, this service recognises "
    "that sometimes you need a quick answer from an authoritative source.\n\n"
    "Contact the Legislation Hotline at LegislationHotline@saiglobal.com."
)

_BOLD_RE = re.compile(r"^\*\*(.+)\*\*$")
_ITALIC_SOURCE_RE = re.compile(r"^\*\((.+)\)\*$")


# ---------------------------------------------------------------------------
# .docx helpers
# ---------------------------------------------------------------------------


def _add_horizontal_rule(doc: Document) -> None:
    """Insert a paragraph with a bottom border acting as a horizontal rule."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()  # noqa: SLF001 — internal docx access required
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_entry(doc: Document, draft_text: str) -> None:
    """Parse and append one draft entry's formatted text to the document."""
    pending_bullets: list[str] = []

    def _flush_bullets() -> None:
        for item in pending_bullets:
            doc.add_paragraph(item, style="List Bullet")
        pending_bullets.clear()

    for raw_line in draft_text.strip().split("\n"):
        line = raw_line.strip()
        if not line:
            _flush_bullets()
            continue

        bold_m = _BOLD_RE.match(line)
        italic_m = _ITALIC_SOURCE_RE.match(line)

        if bold_m:
            _flush_bullets()
            p = doc.add_paragraph()
            p.add_run(bold_m.group(1)).bold = True
        elif italic_m:
            _flush_bullets()
            p = doc.add_paragraph()
            p.add_run(italic_m.group(1)).italic = True
        elif line.startswith("- "):
            pending_bullets.append(line[2:])
        else:
            _flush_bullets()
            doc.add_paragraph(line)

    _flush_bullets()


def _group_by_jurisdiction(
    entries: list[AssembledEntry],
    jurisdiction_order: list[str],
) -> list[tuple[str, list[AssembledEntry]]]:
    """Return [(jurisdiction, entries), ...] in config order."""
    by_juris: dict[str, list[AssembledEntry]] = {}
    for e in entries:
        by_juris.setdefault(e.jurisdiction, []).append(e)

    ordered: list[tuple[str, list[AssembledEntry]]] = []
    for j in jurisdiction_order:
        if j in by_juris:
            ordered.append((j, by_juris[j]))
    for j, ents in by_juris.items():
        if j not in jurisdiction_order:
            ordered.append((j, ents))
    return ordered


# ---------------------------------------------------------------------------
# .docx builder
# ---------------------------------------------------------------------------


def _build_docx(newsfeed: AssembledNewsfeed, run_date: str) -> Document:
    doc = _make_docx()

    # Title
    doc.add_heading(_NEWSFEED_TITLE, level=0)
    subtitle = doc.add_paragraph(f"Issue date: {run_date}")
    subtitle.runs[0].italic = True
    _add_horizontal_rule(doc)

    # Index
    doc.add_heading("Index", level=2)
    for i, section in enumerate(newsfeed.sections, 1):
        doc.add_paragraph(f"{i}. {section.section}", style="List Number")
    upcoming_num = len(newsfeed.sections) + 1
    doc.add_paragraph(f"{upcoming_num}. Upcoming Dates", style="List Number")
    doc.add_paragraph(f"{upcoming_num + 1}. Previous Water Newsfeeds", style="List Number")
    _add_horizontal_rule(doc)

    # Legislation Hotline boilerplate
    doc.add_heading("Do You Have A Question About Legislation?", level=2)
    for para in _LEGISLATION_HOTLINE_BODY.split("\n\n"):
        doc.add_paragraph(para)
    _add_horizontal_rule(doc)

    settings = get_settings()
    jurisdiction_order = settings.jurisdictions

    # Content sections
    for idx, section_data in enumerate(newsfeed.sections, 1):
        doc.add_heading(f"{idx}. {section_data.section}", level=1)

        for juris, entries in _group_by_jurisdiction(section_data.entries, jurisdiction_order):
            doc.add_heading(f"{juris} - {section_data.section}", level=2)
            for entry in entries:
                _add_entry(doc, entry.draft_text)
                if entry.source_url:
                    url_para = doc.add_paragraph()
                    url_para.add_run(entry.source_url).italic = True

        _add_horizontal_rule(doc)

    # Upcoming Dates
    doc.add_heading(f"{upcoming_num}. Upcoming Dates", level=1)
    if newsfeed.upcoming_dates:
        doc.add_paragraph(newsfeed.upcoming_dates)
    else:
        doc.add_paragraph("(No upcoming dates recorded for this issue.)")
    _add_horizontal_rule(doc)

    # Previous Newsfeeds
    doc.add_heading(f"{upcoming_num + 1}. Previous Water Newsfeeds", level=1)
    if newsfeed.previous_newsfeeds:
        doc.add_paragraph(newsfeed.previous_newsfeeds)
    else:
        doc.add_paragraph("(See newsfeed archive for previous issues.)")
    _add_horizontal_rule(doc)

    # Copyright footer
    footer_para = doc.add_paragraph()
    footer_para.add_run(newsfeed.copyright_footer).italic = True

    return doc


# ---------------------------------------------------------------------------
# .md builder
# ---------------------------------------------------------------------------


def _build_md(newsfeed: AssembledNewsfeed, run_date: str) -> str:
    lines: list[str] = [
        f"# {_NEWSFEED_TITLE}",
        "",
        f"*Issue date: {run_date}*",
        "",
        "---",
        "",
        "## Index",
        "",
    ]

    for i, section in enumerate(newsfeed.sections, 1):
        lines.append(f"{i}. {section.section}")

    upcoming_num = len(newsfeed.sections) + 1
    lines += [
        f"{upcoming_num}. Upcoming Dates",
        f"{upcoming_num + 1}. Previous Water Newsfeeds",
        "",
        "---",
        "",
    ]

    settings = get_settings()
    jurisdiction_order = settings.jurisdictions

    for idx, section_data in enumerate(newsfeed.sections, 1):
        lines += [f"## {idx}. {section_data.section}", ""]
        for juris, entries in _group_by_jurisdiction(section_data.entries, jurisdiction_order):
            lines += [f"### {juris} - {section_data.section}", ""]
            for entry in entries:
                lines += [entry.draft_text.strip(), ""]
                if entry.source_url:
                    lines += [f"*{entry.source_url}*", ""]
        lines += ["---", ""]

    lines += [
        f"## {upcoming_num}. Upcoming Dates",
        "",
        newsfeed.upcoming_dates or "(No upcoming dates recorded for this issue.)",
        "",
        "---",
        "",
        f"## {upcoming_num + 1}. Previous Water Newsfeeds",
        "",
        newsfeed.previous_newsfeeds or "(See newsfeed archive for previous issues.)",
        "",
        "---",
        "",
        f"*{newsfeed.copyright_footer}*",
    ]

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def run(
    run_date: str | None = None,
    output_dir: Path | None = None,
    payload_path: Path | None = None,
) -> dict[str, int]:
    """Render the assembled newsfeed to .docx and .md.

    Returns ``{"entries": N, "sections": M}``.
    """
    settings = get_settings()
    today = run_date or date.today().isoformat()
    out_dir = output_dir or (settings.paths.outputs / today)
    src_path = payload_path or (out_dir / "horizon_payload.json")

    if not src_path.exists():
        raise FileNotFoundError(f"horizon_payload.json not found: {src_path}")

    newsfeed = AssembledNewsfeed.model_validate_json(src_path.read_text(encoding="utf-8"))

    total_entries = sum(len(s.entries) for s in newsfeed.sections)
    if total_entries == 0:
        logger.warning("render: no entries in payload — nothing to render")
        return {"entries": 0, "sections": 0}

    out_dir.mkdir(parents=True, exist_ok=True)

    # .docx
    docx_path = out_dir / "Water_Newsfeed.docx"
    doc = _build_docx(newsfeed, today)
    try:
        doc.save(str(docx_path))
    except PermissionError:
        # File is locked (e.g. open in Word) — write to a timestamped fallback
        from datetime import datetime  # noqa: PLC0415

        ts = datetime.now().strftime("%H%M%S")
        docx_path = out_dir / f"Water_Newsfeed_{ts}.docx"
        doc.save(str(docx_path))
        logger.warning("render: original .docx locked, wrote fallback", path=str(docx_path))
    logger.info("render: wrote .docx", path=str(docx_path))

    # .md
    md_path = out_dir / "Water_Newsfeed.md"
    md_text = _build_md(newsfeed, today)
    md_path.write_text(md_text, encoding="utf-8")
    logger.info("render: wrote .md", path=str(md_path))

    counts = {"entries": total_entries, "sections": len(newsfeed.sections)}
    logger.info("render complete", **counts)
    return counts


if __name__ == "__main__":
    date_arg = sys.argv[2] if len(sys.argv) > 2 and sys.argv[1] == "--date" else None
    run(run_date=date_arg)

"""Unit tests for Stage 11 — render to .docx and .md."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document as make_document
from docx.document import Document as DocxDocument

from newsfeed.schemas import AssembledEntry, AssembledNewsfeed, AssembledSection
from newsfeed.stage11_render import (
    _add_entry,
    _build_docx,
    _build_md,
    _group_by_jurisdiction,
    run,
)


def _make_newsfeed(
    sections: list[tuple[str, list[tuple[str, str]]]] | None = None,
) -> AssembledNewsfeed:
    """Build a minimal AssembledNewsfeed for testing.

    ``sections`` is a list of (section_name, [(jurisdiction, draft_text), ...]).
    """
    if sections is None:
        sections = [
            (
                "Water Reform",
                [
                    (
                        "Victoria",
                        "**Water Reform Update**\n"
                        "DEECA announced reforms on 15 April 2024.\n"
                        "*(Source: DEECA)*",
                    )
                ],
            )
        ]

    assembled_sections: list[AssembledSection] = []
    entry_id = 1
    for section_name, juris_entries in sections:
        entries: list[AssembledEntry] = []
        for jurisdiction, draft_text in juris_entries:
            entries.append(
                AssembledEntry(
                    candidate_id=entry_id,
                    draft_id=entry_id,
                    section=section_name,
                    jurisdiction=jurisdiction,
                    content_type="media_release",
                    draft_text=draft_text,
                )
            )
            entry_id += 1
        assembled_sections.append(AssembledSection(section=section_name, entries=entries))

    return AssembledNewsfeed(run_date="2024-04-22", sections=assembled_sections)


# ---------------------------------------------------------------------------
# _group_by_jurisdiction
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_group_by_jurisdiction_respects_order() -> None:
    entries = [
        AssembledEntry(
            candidate_id=1,
            draft_id=1,
            section="s",
            jurisdiction="Victoria",
            content_type="media_release",
            draft_text="x",
        ),
        AssembledEntry(
            candidate_id=2,
            draft_id=2,
            section="s",
            jurisdiction="National",
            content_type="media_release",
            draft_text="y",
        ),
    ]
    order = ["National", "Victoria", "New South Wales"]
    grouped = _group_by_jurisdiction(entries, order)
    jurisdictions = [j for j, _ in grouped]
    assert jurisdictions == ["National", "Victoria"]


@pytest.mark.unit
def test_group_by_jurisdiction_unknown_at_end() -> None:
    entries = [
        AssembledEntry(
            candidate_id=1,
            draft_id=1,
            section="s",
            jurisdiction="Atlantis",
            content_type="media_release",
            draft_text="x",
        ),
        AssembledEntry(
            candidate_id=2,
            draft_id=2,
            section="s",
            jurisdiction="National",
            content_type="media_release",
            draft_text="y",
        ),
    ]
    order = ["National", "Victoria"]
    grouped = _group_by_jurisdiction(entries, order)
    jurisdictions = [j for j, _ in grouped]
    assert jurisdictions[0] == "National"
    assert "Atlantis" in jurisdictions


# ---------------------------------------------------------------------------
# _add_entry (docx paragraph rendering)
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_add_entry_bold_title() -> None:
    doc = make_document()
    _add_entry(doc, "**Water Reform Update**\nBody text here.\n*(Source: DEECA)*")
    paragraphs = doc.paragraphs
    bold_paras = [p for p in paragraphs if any(r.bold for r in p.runs)]
    assert len(bold_paras) >= 1
    assert bold_paras[0].runs[0].text == "Water Reform Update"


@pytest.mark.unit
def test_add_entry_italic_source() -> None:
    doc = make_document()
    _add_entry(doc, "**Title**\nBody.\n*(Source: DEECA)*")
    paragraphs = doc.paragraphs
    italic_paras = [p for p in paragraphs if any(r.italic for r in p.runs)]
    assert len(italic_paras) >= 1
    assert "Source: DEECA" in italic_paras[0].text


@pytest.mark.unit
def test_add_entry_bullet_list() -> None:
    doc = make_document()
    draft = "**Agency Releases**\n- First item (1 Jan 2024);\n- Second item (2 Jan 2024).\n*(Source: Agency)*"
    _add_entry(doc, draft)
    # Bullet list paragraphs should exist
    bullet_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert len(bullet_paras) == 2


@pytest.mark.unit
def test_add_entry_handles_empty_lines() -> None:
    doc = make_document()
    _add_entry(doc, "**Title**\n\nBody after blank line.\n*(Source: Agency)*")
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert "Title" in texts[0]
    assert "Body after blank line." in texts


# ---------------------------------------------------------------------------
# _build_md
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_build_md_contains_title() -> None:
    newsfeed = _make_newsfeed()
    md = _build_md(newsfeed, "2024-04-22")
    assert "# Water Newsfeed" in md


@pytest.mark.unit
def test_build_md_contains_section_heading() -> None:
    newsfeed = _make_newsfeed()
    md = _build_md(newsfeed, "2024-04-22")
    assert "Water Reform" in md


@pytest.mark.unit
def test_build_md_contains_jurisdiction_subheading() -> None:
    newsfeed = _make_newsfeed()
    md = _build_md(newsfeed, "2024-04-22")
    assert "Victoria - Water Reform" in md


@pytest.mark.unit
def test_build_md_contains_draft_text() -> None:
    newsfeed = _make_newsfeed()
    md = _build_md(newsfeed, "2024-04-22")
    assert "Water Reform Update" in md
    assert "DEECA announced reforms" in md


@pytest.mark.unit
def test_build_md_contains_copyright_footer() -> None:
    newsfeed = _make_newsfeed()
    md = _build_md(newsfeed, "2024-04-22")
    assert "Legalwise" in md or "rights reserved" in md


@pytest.mark.unit
def test_build_md_contains_index() -> None:
    newsfeed = _make_newsfeed()
    md = _build_md(newsfeed, "2024-04-22")
    assert "Index" in md
    assert "Upcoming Dates" in md
    assert "Previous Water Newsfeeds" in md


# ---------------------------------------------------------------------------
# _build_docx
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_build_docx_returns_document() -> None:
    newsfeed = _make_newsfeed()
    doc = _build_docx(newsfeed, "2024-04-22")
    assert isinstance(doc, DocxDocument)


@pytest.mark.unit
def test_build_docx_contains_section_heading() -> None:
    newsfeed = _make_newsfeed()
    doc = _build_docx(newsfeed, "2024-04-22")
    heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("Water Reform" in t for t in heading_texts)


@pytest.mark.unit
def test_build_docx_multiple_sections() -> None:
    newsfeed = _make_newsfeed(
        sections=[
            ("Water Reform", [("Victoria", "**Title A**\n*(Source: DEECA)*")]),
            ("Conservation", [("National", "**Title B**\n*(Source: DCCEEW)*")]),
        ]
    )
    doc = _build_docx(newsfeed, "2024-04-22")
    heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert any("Conservation" in t for t in heading_texts)
    assert any("Water Reform" in t for t in heading_texts)


# ---------------------------------------------------------------------------
# run()
# ---------------------------------------------------------------------------


@pytest.mark.unit
def test_run_creates_docx_and_md(tmp_path: Path) -> None:
    newsfeed = _make_newsfeed()
    payload_path = tmp_path / "horizon_payload.json"
    payload_path.write_text(newsfeed.model_dump_json(), encoding="utf-8")

    counts = run(run_date="2024-04-22", output_dir=tmp_path, payload_path=payload_path)

    assert counts["entries"] == 1
    assert counts["sections"] == 1
    assert (tmp_path / "Water_Newsfeed.docx").exists()
    assert (tmp_path / "Water_Newsfeed.md").exists()


@pytest.mark.unit
def test_run_raises_on_missing_payload(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        run(run_date="2024-04-22", output_dir=tmp_path, payload_path=tmp_path / "nope.json")


@pytest.mark.unit
def test_run_returns_zeros_for_empty_payload(tmp_path: Path) -> None:
    newsfeed = AssembledNewsfeed(run_date="2024-04-22", sections=[])
    payload_path = tmp_path / "horizon_payload.json"
    payload_path.write_text(newsfeed.model_dump_json(), encoding="utf-8")

    counts = run(run_date="2024-04-22", output_dir=tmp_path, payload_path=payload_path)

    assert counts == {"entries": 0, "sections": 0}


@pytest.mark.unit
def test_run_md_content_matches_draft_text(tmp_path: Path) -> None:
    newsfeed = _make_newsfeed()
    payload_path = tmp_path / "horizon_payload.json"
    payload_path.write_text(newsfeed.model_dump_json(), encoding="utf-8")

    run(run_date="2024-04-22", output_dir=tmp_path, payload_path=payload_path)

    md = (tmp_path / "Water_Newsfeed.md").read_text(encoding="utf-8")
    assert "Water Reform Update" in md
    assert "DEECA announced reforms" in md
